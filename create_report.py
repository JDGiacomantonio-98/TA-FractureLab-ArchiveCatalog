def main():
    return generate_report(load_data())


def load_data(filepath=None):
    ''' Retrieves and pack data from the excel database of archive resources '''

    from xlrd import open_workbook

    db = open_workbook("Lab-of-Fracture database.xlsx" if filepath is None else filepath).sheet_by_index(0)

    keys = set()
    for r in range(1, db.nrows-1):
        keys.add(db.cell_value(r, 0))
        archive = dict.fromkeys(keys, {})

    keys = []
    for h in range(2, db.ncols):
        keys.append(db.cell_value(0, h).lower())
    
    for f in archive.keys():
        index = {}
        for r in range(1, db.nrows-1):
            if db.cell_value(r, 0) != f:
                continue
            document = {}
            for i, h in enumerate(range(2, db.ncols)):
                document[keys[i]] = str(db.cell_value(r, h)) if keys[i] not in ("pub_year", "vol_n", "paper_n") else str(db.cell_value(r, h))[:len(str(db.cell_value(r, h)))-2] if ".0" in str(db.cell_value(r, h)) else str(db.cell_value(r, h))
            index[int(db.cell_value(r, 1))] = document
        archive[f] = index
    
    return archive


def generate_report(archive):
    ''' Generates the requested Word file accordingly to all the authors formatting rules '''

    from os import path, mkdir, getcwd
    from datetime import datetime as dt

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not path.isdir(f'{getcwd()}\\reports'):
        mkdir(f'{getcwd()}\\reports')

    for folder, index in archive.items():
        doc = Document()
        doc.styles['Normal'].font.name = "Times New Roman"
        doc.add_paragraph(f'FALDONE : {folder}')
        for p in index.values():
            par = doc.add_paragraph(f'{p["authors"]}: "{p["title"]}"', style="List Number")
            if p["magazine_name"]:
                par.add_run(f', {p["magazine_name"]}').italic = True
            if p["vol_n"]:
                par.add_run(f', Vol. {p["vol_n"]}')
            if p["paper_n"]:
                par.add_run(f', Paper N. {p["paper_n"]}')
            if p["pub_year"]:
                par.add_run(f'{", " if not p["vol_n"] else " "}({str(p["pub_year"])})')
            if p["pp_interval"]:
                par.add_run(f', {p["pp_interval"]}')
            par.add_run('.')
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.save(f'{getcwd()}\\reports\\CATALOGO {folder.upper()}.docx')

main()